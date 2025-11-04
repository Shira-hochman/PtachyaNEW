# generate_pdf_from_docx.py
import json
import sys
import os
import io
import base64
import subprocess
import shutil  # ⭐️ חובה לפתרון עם soffice.exe
from mailmerge import MailMerge
from datetime import datetime
from docx import Document
from docx.shared import Inches

# נתיב קבוע לספרייה זמנית (כדי לאסוף את הפלט של LibreOffice)
TEMP_DIR = os.path.join(os.path.expanduser('~'), 'ptachya_temp_forms')
os.makedirs(TEMP_DIR, exist_ok=True)


# ----------------- פונקציות עזר -----------------

def format_date(date_str: str) -> str:
    """ממיר תאריך YYYY-MM-DD ל-DD/MM/YYYY."""
    try:
        if not date_str: return "_______"
        date_part = date_str.split('T')[0]
        return datetime.strptime(date_part, '%Y-%m-%d').strftime('%d/%m/%Y')
    except:
        return "_______"


def get_mark(form_value: str, check_val: str) -> str:
    """מחזיר 'X' אם הערך תואם לסימון, אחרת מחרוזת ריקה."""
    return 'X' if form_value == check_val else ''


def InsertSignaturePlaceholder(doc, placeholder_text, base64_data):
    """מוצא ומחליף טקסט Placeholder בתמונת Base64."""
    if not base64_data:
        for p in doc.paragraphs:
            if placeholder_text in p.text:
                p.text = p.text.replace(placeholder_text, '_______ (לא נחתם) _______')
        return

    base64WithoutPrefix = base64_data.split(',')[-1]
    try:
        image_bytes = base64.b64decode(base64WithoutPrefix)
    except Exception as e:
        print(f"Error decoding Base64 for {placeholder_text}: {e}", file=sys.stderr)
        return

    for p in doc.paragraphs:
        if placeholder_text in p.text:
            p.text = p.text.replace(placeholder_text, '')

            image_stream = io.BytesIO(image_bytes)
            p.add_run().add_picture(image_stream, width=Inches(1.8), height=Inches(0.5))
            break


# ----------------- הפונקציה המרכזית: מילוי והמרה -----------------

# ⭐️ הפונקציה מקבלת את הנתיב הגמיש מה-C#
def fill_and_convert_to_pdf(data: dict, output_pdf_path: str, template_path: str, libre_office_path: str):
    """ממלא את התבנית DOCX, מטמיע חתימות, וממיר ל-PDF באמצעות soffice.exe."""

    child = data.get('childDetails', {})
    facility = data.get('facilityDetails', {})
    parent1 = data.get('parent1', {})
    parent2 = data.get('parent2', {})

    # ⭐️ מיפוי הנתונים לשדות ה-Mail Merge
    merge_data = {
        'FormDate': format_date(data.get('formDate', '')),
        'ChildFirstName': child.get('childFirstName', ''),
        'ChildLastName': child.get('childLastName', ''),
        'IDNumber': child.get('childId', ''),
        'ChildDateOfBirth': format_date(child.get('childDob', '')),
        'ChildAddress': child.get('childAddress', ''),
        'ProgramProvider': data.get('programProvider', ''),
        'SelfParticipation': str(data.get('monthlySelfParticipation', '')),
        'FacilityName': facility.get('facilityName', ''),
        'ManagerName': facility.get('facilityManagerName', ''),
        'FacilityAddress': facility.get('facilityAddress', ''),
        'FacilityPhone': facility.get('facilityPhone', ''),
        'Parent1Name': parent1.get('name', ''),
        'Parent1Phone': parent1.get('phone', ''),
        'Parent2Name': parent2.get('name', ''),
        'Parent2Phone': parent2.get('phone', ''),

        # שדות סימון (חייבים להתאים למה שמופיע ב-DOCX)
        'Mark_Gan': get_mark(data.get('programFramework', ''), 'גן תקשורתי'),
        'Mark_Merkaz': get_mark(data.get('programFramework', ''), 'מרכז טיפולי'),
        'Mark_Local': get_mark(facility.get('facilityOwnership', ''), 'רשות מקומית'),
        'Mark_Private': get_mark(facility.get('facilityOwnership', ''), 'בעלות פרטית עמותה'),
        'Mark_Declaration': 'X' if data.get('noOtherProgramDeclaration') else '',
    }

    # 1. מילוי התבנית DOCX
    temp_docx_base_name = os.path.basename(output_pdf_path).replace('.pdf', '')
    temp_docx_path = os.path.join(TEMP_DIR, f"filled_temp_{temp_docx_base_name}.docx")

    with MailMerge(template_path) as document:
        document.merge(**merge_data)
        document.write(temp_docx_path)

    # 2. הטמעת חתימות Base64
    doc = Document(temp_docx_path)
    # ⭐️ Parent1Signature_PH חייב להיות ה-Placeholder שהגדרת בתבנית
    InsertSignaturePlaceholder(doc, 'Parent1Signature_PH', parent1.get('signature', ''))
    doc.save(temp_docx_path)

    # 3. המרת DOCX ל-PDF באמצעות SOFFICE (LibreOffice)
    expected_pdf_name = os.path.basename(temp_docx_path).replace('.docx', '.pdf')
    generated_pdf_path = os.path.join(TEMP_DIR, expected_pdf_name)

    try:
        subprocess.run(
            [
                libre_office_path,  # ⭐️ הנתיב הגמיש שהגיע מה-appsettings ⭐️
                '--headless',
                '--convert-to', 'pdf',
                temp_docx_path,
                '--outdir', TEMP_DIR
            ],
            capture_output=True, text=True, check=True, timeout=60
        )

        # 4. העתקת הקובץ למיקום שה-C# מצפה לו
        shutil.copyfile(generated_pdf_path, output_pdf_path)

    except subprocess.CalledProcessError as e:
        # שגיאה קשורה ל-LibreOffice עצמו (למשל, קובץ DOCX פגום)
        raise Exception(f"PDF Conversion failed (soffice). Error: {e.stderr}")
    except FileNotFoundError:
        # שגיאה זו תתרחש אם הנתיב שהועבר ל-libre_office_path שגוי
        raise Exception(
            f"LibreOffice command not found. Please check 'LibreOfficeExecutablePath' in appsettings: {libre_office_path}")

    finally:
        # 5. ניקוי קבצים זמניים
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        if os.path.exists(generated_pdf_path):
            os.remove(generated_pdf_path)


# ----------------- הפונקציה הראשית להפעלה -----------------

def main():
    try:
        input_json_data = sys.stdin.read()
        data_wrapper = json.loads(input_json_data)

        output_path = data_wrapper.get('output_pdf_path')
        template_path = data_wrapper.get('template_path')
        form_data = data_wrapper.get('form_data', {})
        libre_office_path = data_wrapper.get('libre_office_path')

        if not output_path or not template_path or not libre_office_path:
            raise ValueError("נתיב שמירה, נתיב תבנית או נתיב LibreOffice חסרים.")

        fill_and_convert_to_pdf(form_data, output_path, template_path, libre_office_path)

        print(json.dumps({'status': 'success', 'path': output_path}))

    except Exception as e:
        print(json.dumps({'status': 'error', 'message': str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()