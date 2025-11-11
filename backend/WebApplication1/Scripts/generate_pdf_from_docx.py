import json
import sys
import os
import io
import base64
import subprocess
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# נתיב קבוע לספרייה זמנית
TEMP_DIR = os.path.join(os.path.expanduser('~'), 'ptachya_temp_forms')
os.makedirs(TEMP_DIR, exist_ok=True)

FOUR_SPACES = '    '
NINE_UNDERSCORES = '_________'


# ----------------- פונקציות עזר -----------------

def format_date(date_str: str) -> str:
    try:
        if not date_str:
            return "_______"
        date_part = date_str.split('T')[0]
        return datetime.strptime(date_part, '%Y-%m-%d').strftime('%d/%m/%Y')
    except:
        return "_______"


def get_mark(form_value: str, check_val: str) -> str:
    return 'X' if str(form_value) == str(check_val) else ' '


def InsertSignaturePlaceholder(doc, placeholder_text, base64_data):
    if not base64_data:
        for p in doc.paragraphs:
            if placeholder_text in p.text:
                p.text = p.text.replace(placeholder_text, '_______ (לא נחתם) _______')
        return

    base64WithoutPrefix = base64_data.split(',')[-1]
    try:
        image_bytes = base64.b64decode(base64WithoutPrefix)
    except Exception as e:
        print(f"Error decoding Base64 for {placeholder_text}: {e}", file=sys.stderr)
        return

    for p in doc.paragraphs:
        if placeholder_text in p.text:
            p.text = p.text.replace(placeholder_text, '')
            image_stream = io.BytesIO(image_bytes)
            p.add_run().add_picture(image_stream, width=Inches(1.8), height=Inches(0.5))
            break


# ----------------- ⭐️ פונקציית ההחלפה הכללית (לשימור תוכן) ⭐️ -----------------

def replace_and_style_stable(paragraph, key, value, special_handling=False):
    if key not in paragraph.text:
        return False

    original_font = None
    for run in paragraph.runs:
        if key in run.text:
            original_font = run.font
            break

    # קביעת ערך התוכן למילוי (עם קווים תחתונים אם ריק)
    is_mark = (value == 'X' or value == ' ')
    if is_mark:
        content_value = str(value)
        underline_value = False
    elif value:
        content_value = f" {str(value)} "
        underline_value = True
    else:
        content_value = NINE_UNDERSCORES
        underline_value = False # אם זה קווים, לא צריך קו תחתון פונטי נוסף
        if special_handling: # עבור שדות הילד, נשתמש בקו ארוך יותר במקרה של ריק
             content_value = " " + "_" * 10 + " "

    original_text = paragraph.text
    temp_placeholder = f"@@@TEMP_PH_{key}@@@"
    
    # 1. החלפת הטקסט הפנימי במחרוזת זמנית
    if key in original_text:
        original_text = original_text.replace(key, temp_placeholder)
        
        # 2. ניקוי הפסקה (נחוץ כדי להבטיח שהעיצוב נבנה מחדש נכון)
        paragraph.clear()
        parts = original_text.split(temp_placeholder)

        # 3. הוספת הריצות
        for i, part in enumerate(parts):
            paragraph.add_run(part)
            
            if i < len(parts) - 1:
                if not underline_value:
                    # מילוי ב-X או קווים קבועים (NINE_UNDERSCORES)
                    paragraph.add_run(content_value)
                else:
                    # מילוי בערך אמיתי עם קו תחתון פונטי
                    new_run = paragraph.add_run(content_value)
                    if original_font:
                        new_run.font.name = original_font.name
                        new_run.font.size = original_font.size
                        new_run.font.bold = original_font.bold
                    new_run.font.underline = True 
                
                # טיפול מיוחד לשדות הילד עם רווחים
                if special_handling and i < len(parts) - 1:
                    paragraph.add_run(FOUR_SPACES)
    return True


# ----------------- ⭐️ הפונקציה המרכזית ⭐️ -----------------

def fill_and_convert_to_pdf(data: dict, output_pdf_path: str, template_path: str, libre_office_path: str):
    child = data.get('childDetails', {})
    facility = data.get('facilityDetails', {})
    parent1 = data.get('parent1', {})
    parent2 = data.get('parent2', {})

    # מחזירים את מפתחות הילד למפה
    replace_map = {
        '{{FormDate}}': format_date(data.get('formDate', '')),
        '{{ProgramProvider}}': data.get('programProvider', ''),
        '{{SelfParticipation}}': str(data.get('monthlySelfParticipation', '')),
        '{{ChildFirstName}}': child.get('childFirstName', ''),
        '{{ChildLastName}}': child.get('childLastName', ''),
        '{{ChildID}}': child.get('childId', ''),
        '{{ChildDateOfBirth}}': format_date(child.get('childDob', '')),
        '{{ChildAddress}}': child.get('childAddress', ''),
        '{{FacilityName}}': facility.get('facilityName', ''),
        '{{ManagerName}}': facility.get('facilityManagerName', ''),
        '{{FacilityAddress}}': facility.get('facilityAddress', ''),
        '{{FacilityPhone}}': facility.get('facilityPhone', ''),
        '{{Parent1Name}}': parent1.get('name', ''),
        '{{Parent1Phone}}': parent1.get('phone', ''),
        '{{Parent2Name}}': parent2.get('name', ''),
        '{{Parent2Phone}}': parent2.get('phone', ''),
        '{{Mark_Gan}}': get_mark(data.get('programFramework', ''), 'גן תקשורתי'),
        '{{Mark_Merkaz}}': get_mark(data.get('programFramework', ''), 'מרכז טיפולי'),
        '{{Mark_Local}}': get_mark(facility.get('facilityOwnership', ''), 'רשות מקומית'),
        '{{Mark_Private}}': get_mark(facility.get('facilityOwnership', ''), 'בעלות פרטית עמותה'),
        '{{Mark_Declaration}}': get_mark(data.get('noOtherProgramDeclaration'), True),
    }

    # מפתחות מיוחדים לשדות הילד הבעייתיים
    child_special_keys = {'{{ChildFirstName}}', '{{ChildLastName}}', '{{ChildID}}', '{{ChildDateOfBirth}}'}
    
    temp_docx_base_name = os.path.basename(output_pdf_path).replace('.pdf', '')
    temp_docx_path = os.path.join(TEMP_DIR, f"filled_temp_{temp_docx_base_name}.docx")

    # לוודא שהתבנית אכן קיימת
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found at: {template_path}")

    shutil.copyfile(template_path, temp_docx_path)
    doc = Document(temp_docx_path)

    # ⭐️ שלב 1: החלפת השדות ⭐️
    def process_element(element):
        for paragraph in element.paragraphs:
            # 1. טיפול בשדות הילד המופיעים בפסקה אחת (שיטת השימור)
            is_child_paragraph = any(key in paragraph.text for key in child_special_keys)
            
            if is_child_paragraph:
                for key in child_special_keys:
                    # שימוש ב-replace_and_style_stable
                    replace_and_style_stable(paragraph, key, replace_map.get(key, ''), special_handling=True)
            
            # 2. טיפול בכל שאר המפתחות (כולל שדות הילד שכבר טופלו בשלב 1)
            for key, value in replace_map.items():
                if key not in child_special_keys or not is_child_paragraph:
                    # שימוש ב-replace_and_style המקורית
                    replace_and_style_stable(paragraph, key, value, special_handling=False)


    # מעבר על כל הפסקאות והטבלאות במסמך
    process_element(doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_element(cell)

    # טיפול בחתימות
    InsertSignaturePlaceholder(doc, '{{Parent1Signature_PH}}', parent1.get('signature', ''))
    InsertSignaturePlaceholder(doc, '{{Parent2Signature_PH}}', parent2.get('signature', ''))

    doc.save(temp_docx_path)

    expected_pdf_name = os.path.basename(temp_docx_path).replace('.docx', '.pdf')
    generated_pdf_path = os.path.join(TEMP_DIR, expected_pdf_name)

    try:
        subprocess.run(
            [libre_office_path, '--headless', '--convert-to', 'pdf:writer_pdf_Export',
             temp_docx_path, '--outdir', TEMP_DIR],
            capture_output=True, text=True, check=True, timeout=60
        )
        shutil.copyfile(generated_pdf_path, output_pdf_path)
    finally:
        for path in [temp_docx_path, generated_pdf_path]:
            if os.path.exists(path):
                os.remove(path)


# ----------------- הפונקציה הראשית -----------------

def main():
    try:
        input_json_data = sys.stdin.read()
        data_wrapper = json.loads(input_json_data)
        output_path = data_wrapper.get('output_pdf_path')
        template_path = data_wrapper.get('template_path')
        form_data = data_wrapper.get('form_data', {})
        libre_office_path = data_wrapper.get('libre_office_path')

        if not output_path or not template_path or not libre_office_path:
            raise ValueError("נתיב שמירה, נתיב תבנית או נתיב LibreOffice חסרים.")

        fill_and_convert_to_pdf(form_data, output_path, template_path, libre_office_path)
        print(json.dumps({'status': 'success', 'path': output_path}))
    except Exception as e:
        print(json.dumps({'status': 'error', 'message': str(e)}), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()